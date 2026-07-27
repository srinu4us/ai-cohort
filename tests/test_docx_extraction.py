from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from main import app

client = TestClient(app)


def test_extract_claims_process_doc_returns_paragraph_text(tmp_path):
    doc_path = tmp_path / "claims_process.docx"
    document = Document()
    document.add_paragraph("Claim intake workflow")
    document.add_paragraph("Review coverage details")
    document.save(doc_path)

    response = client.post("/claims-process-doc/extract", params={"docx_path": str(doc_path)})

    assert response.status_code == 200
    payload = response.json()
    assert payload["docx_path"] == str(doc_path)
    assert payload["paragraphs"] == ["Claim intake workflow", "Review coverage details"]
